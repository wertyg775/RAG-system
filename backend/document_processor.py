"""Document parsing for various file formats."""
import io
from pathlib import Path
from typing import Dict, Any, List
import PyPDF2
import pdfplumber
import docx
import markdown


class DocumentMetadata:
    """Metadata for a parsed document."""
    
    def __init__(
        self,
        filename: str,
        file_type: str,
        total_pages: int = 1,
        author: str = "",
        title: str = ""
    ):
        self.filename = filename
        self.file_type = file_type
        self.total_pages = total_pages
        self.author = author
        self.title = title
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "author": self.author,
            "title": self.title
        }


class ParsedDocument:
    """Represents a parsed document with text and metadata."""
    
    def __init__(self, text: str, metadata: DocumentMetadata, page_texts: List[str] = None):
        self.text = text
        self.metadata = metadata
        self.page_texts = page_texts or [text]  # Text split by pages if available


class PDFParser:
    """Parse PDF documents."""
    
    @staticmethod
    def parse(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF and extract text with page information."""
        text_parts = []
        page_texts = []
        
        try:
            # Use pdfplumber for better text extraction
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                total_pages = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    page_texts.append(page_text)
                    text_parts.append(page_text)
                
                # Try to extract metadata
                metadata_dict = pdf.metadata or {}
                author = metadata_dict.get("Author", "")
                title = metadata_dict.get("Title", filename)
        
        except Exception as e:
            # Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            total_pages = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
                text_parts.append(page_text)
            
            info = pdf_reader.metadata or {}
            author = info.get("/Author", "")
            title = info.get("/Title", filename)
        
        full_text = "\n\n".join(text_parts)
        
        metadata = DocumentMetadata(
            filename=filename,
            file_type="pdf",
            total_pages=total_pages,
            author=author,
            title=title
        )
        
        return ParsedDocument(full_text, metadata, page_texts)


class DOCXParser:
    """Parse DOCX documents."""
    
    @staticmethod
    def parse(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX and extract text."""
        doc = docx.Document(io.BytesIO(file_content))
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        # Try to get metadata
        core_properties = doc.core_properties
        author = core_properties.author or ""
        title = core_properties.title or filename
        
        metadata = DocumentMetadata(
            filename=filename,
            file_type="docx",
            total_pages=1,  # DOCX doesn't have strict pages
            author=author,
            title=title
        )
        
        return ParsedDocument(full_text, metadata)


class MarkdownParser:
    """Parse Markdown documents."""
    
    @staticmethod
    def parse(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse Markdown file."""
        text = file_content.decode('utf-8')
        
        # Convert markdown to HTML for better processing (optional)
        # For now, we'll keep the raw markdown as it's already readable
        
        metadata = DocumentMetadata(
            filename=filename,
            file_type="markdown",
            total_pages=1
        )
        
        return ParsedDocument(text, metadata)


class TXTParser:
    """Parse plain text documents."""
    
    @staticmethod
    def parse(file_content: bytes, filename: str) -> ParsedDocument:
        """Parse plain text file."""
        text = file_content.decode('utf-8')
        
        metadata = DocumentMetadata(
            filename=filename,
            file_type="txt",
            total_pages=1
        )
        
        return ParsedDocument(text, metadata)


class DocumentProcessor:
    """Main document processor that routes to appropriate parser."""
    
    PARSERS = {
        '.pdf': PDFParser,
        '.docx': DOCXParser,
        '.doc': DOCXParser,
        '.md': MarkdownParser,
        '.markdown': MarkdownParser,
        '.txt': TXTParser
    }
    
    @classmethod
    def parse(cls, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse document based on file extension."""
        file_ext = Path(filename).suffix.lower()
        
        parser_class = cls.PARSERS.get(file_ext)
        if not parser_class:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        return parser_class.parse(file_content, filename)
    
    @classmethod
    def supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls.PARSERS.keys())
